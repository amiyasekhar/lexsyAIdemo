from __future__ import annotations

import os
import re
import json
import csv
from dataclasses import dataclass
from typing import List, Dict, Any, Optional, Tuple

from dotenv import load_dotenv
load_dotenv()

import pandas as pd
from docx import Document
from pydantic import BaseModel, Field, ValidationError

# OpenAI SDK (Responses API)
from openai import OpenAI  # per official docs  [oai_citation:1‡OpenAI Platform](https://platform.openai.com/docs/overview?lang=python&utm_source=chatgpt.com)


# -----------------------------
# 1) Deterministic field detection
# -----------------------------

BRACKET_TOKEN_RE = re.compile(r"\[[^\[\]\n]{1,120}\]")  # e.g. [Company Name]
BRACKET_BLANK_RE = re.compile(r"\$\[[\s_]{3,}\]|\[[\s_]{3,}\]")  # e.g. $[_____]
UNDERLINE_RE = re.compile(r"_{3,}")  # e.g. ________
SIGNATURE_LABEL_RE = re.compile(r"^\s*(By:|Name:|Title:|Address:|Email:)\s*(.*)$", re.IGNORECASE)

@dataclass
class DetectedField:
    raw_placeholder: str               # e.g. "[Company Name]" or "By:"
    kind: str                          # "bracket_token" | "bracket_blank" | "underline" | "signature_label"
    context_before: str
    context_line: str
    context_after: str
    location_hint: str                 # "paragraph:12" or "table:3 row:1 col:0"
    line_index: int                    # index in flattened lines


def _iter_docx_lines(doc: Document) -> List[Tuple[str, str]]:
    """
    Flatten doc into a list of (text, location_hint) from paragraphs + tables.
    """
    lines: List[Tuple[str, str]] = []

    # Paragraphs
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t:
            lines.append((t, f"paragraph:{i}"))

    # Tables
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell_text = " ".join((cell.text or "").split())
                if cell_text:
                    lines.append((cell_text, f"table:{ti} row:{ri} col:{ci}"))

    return lines


def detect_fields_from_docx(path: str, window: int = 1) -> List[DetectedField]:
    doc = Document(path)
    lines = _iter_docx_lines(doc)

    detected: List[DetectedField] = []

    for idx, (line, loc) in enumerate(lines):
        before = lines[idx - window][0] if idx - window >= 0 else ""
        after = lines[idx + window][0] if idx + window < len(lines) else ""

        # 1) bracket tokens like [Company Name]
        for m in BRACKET_TOKEN_RE.finditer(line):
            detected.append(
                DetectedField(
                    raw_placeholder=m.group(0),
                    kind="bracket_token",
                    context_before=before,
                    context_line=line,
                    context_after=after,
                    location_hint=loc,
                    line_index=idx,
                )
            )

        # 2) bracket blanks like $[_____]
        for m in BRACKET_BLANK_RE.finditer(line):
            detected.append(
                DetectedField(
                    raw_placeholder=m.group(0),
                    kind="bracket_blank",
                    context_before=before,
                    context_line=line,
                    context_after=after,
                    location_hint=loc,
                    line_index=idx,
                )
            )

        # 3) signature labels ("By:", "Name:", etc.) even without explicit underscores
        sig = SIGNATURE_LABEL_RE.match(line)
        if sig:
            label = sig.group(1).strip()
            detected.append(
                DetectedField(
                    raw_placeholder=label,
                    kind="signature_label",
                    context_before=before,
                    context_line=line,
                    context_after=after,
                    location_hint=loc,
                    line_index=idx,
                )
            )

        # 4) explicit underlines ________ that aren’t already inside bracket_blank
        if UNDERLINE_RE.search(line) and not BRACKET_BLANK_RE.search(line):
            detected.append(
                DetectedField(
                    raw_placeholder="__________",
                    kind="underline",
                    context_before=before,
                    context_line=line,
                    context_after=after,
                    location_hint=loc,
                    line_index=idx,
                )
            )

    # De-dupe by (raw_placeholder, line_index, kind)
    uniq = {}
    for d in detected:
        key = (d.raw_placeholder, d.line_index, d.kind)
        uniq[key] = d
    return list(uniq.values())


# -----------------------------
# 2) LLM classification schema
# -----------------------------

class LLMField(BaseModel):
    field_id: str = Field(..., description="Stable dot-notation id, e.g. company.name, safe.date")
    raw_placeholder: str
    label_for_user: str
    who_should_fill: str = Field(..., description="company|counterparty|either|system")
    expected_type: str = Field(..., description="string|date|money|state|email|name|title|address|jurisdiction|other")
    how_to_fill: str = Field(..., description="One sentence instruction to the user.")
    confidence: float = Field(..., ge=0.0, le=1.0)
    evidence_quote: str
    location_hint: str
    ask_user: bool


# -----------------------------
# 3) OpenAI call (Responses API)
# -----------------------------

SYSTEM_INSTRUCTIONS = """You are a document field-extraction engine for legal documents.
You will receive detected placeholders with local context.
Your job: infer what the field should be, who should fill it, and how to ask the user.

Rules:
- Output ONLY valid JSON: an array of objects matching the schema described.
- Do NOT invent actual values. Only describe what should be entered.
- The end-user is NOT the investor / NOT the paying party / NOT giving money.
  So if the placeholder is clearly for the counterparty (e.g., Investor name, Purchase Amount),
  set who_should_fill="counterparty" and ask_user=false.
- Typed signatures only: for signature blocks, ask for printed name/title/email/address where appropriate, not a drawn signature.
- If uncertain, set lower confidence and explain briefly in how_to_fill.
"""

def call_openai_classify(detected: List[DetectedField], model: str = "gpt-5.2") -> List[LLMField]:
    """
    Sends detected placeholders to OpenAI and expects JSON array back.
    """
    client = OpenAI()  # uses OPENAI_API_KEY env var  [oai_citation:2‡OpenAI Platform](https://platform.openai.com/docs/api-reference/introduction?utm_source=chatgpt.com)

    payload_items: List[Dict[str, Any]] = []
    for d in detected:
        payload_items.append({
            "raw_placeholder": d.raw_placeholder,
            "kind": d.kind,
            "location_hint": d.location_hint,
            "context_before": d.context_before,
            "context_line": d.context_line,
            "context_after": d.context_after,
        })

    # We ask for strict JSON.
    user_input = {
        "schema": {
            "field_id": "string",
            "raw_placeholder": "string",
            "label_for_user": "string",
            "who_should_fill": "company|counterparty|either|system",
            "expected_type": "string|date|money|state|email|name|title|address|jurisdiction|other",
            "how_to_fill": "string (1 sentence)",
            "confidence": "number 0..1",
            "evidence_quote": "string (short excerpt from provided context)",
            "location_hint": "string",
            "ask_user": "boolean"
        },
        "placeholders": payload_items
    }

    resp = client.responses.create(
        model=model,
        instructions=SYSTEM_INSTRUCTIONS,
        input=json.dumps(user_input),
        # If your model/account supports it, you can push even harder for JSON:
        # response_format={"type": "json_object"}
    )

    text = resp.output_text.strip()
    # Expect: JSON array
    data = json.loads(text)

    out: List[LLMField] = []
    for obj in data:
        try:
            out.append(LLMField(**obj))
        except ValidationError as e:
            raise RuntimeError(f"LLM output failed schema validation: {e}\nRaw object: {obj}") from e

    return out


# -----------------------------
# 4) Output writers
# -----------------------------

def write_json(fields: List[LLMField], path: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump([f.model_dump() for f in fields], f, ensure_ascii=False, indent=2)

def write_csv(fields: List[LLMField], path: str) -> None:
    rows = [f.model_dump() for f in fields]
    df = pd.DataFrame(rows)
    df.to_csv(path, index=False)

def main():
    import argparse

    parser = argparse.ArgumentParser(description="Extract + classify fillable fields from a .docx legal document.")
    parser.add_argument("docx_path", help="Path to .docx file")
    parser.add_argument("--window", type=int, default=1, help="Context window (lines before/after)")
    parser.add_argument("--model", type=str, default="gpt-5.2", help="OpenAI model name")
    parser.add_argument("--out_json", type=str, default="fields.json", help="Output JSON path")
    parser.add_argument("--out_csv", type=str, default="fields.csv", help="Output CSV path")
    parser.add_argument("--no_llm", action="store_true", help="Only detect fields, do not call LLM (debug mode)")
    args = parser.parse_args()

    detected = detect_fields_from_docx(args.docx_path, window=args.window)
    print(f"Detected {len(detected)} candidate placeholders/fields.")

    if args.no_llm:
        # Dump raw detections for inspection
        raw_path = "detections.json"
        with open(raw_path, "w", encoding="utf-8") as f:
            json.dump([d.__dict__ for d in detected], f, ensure_ascii=False, indent=2)
        print(f"Wrote raw detections to {raw_path}")
        return

    fields = call_openai_classify(detected, model=args.model)
    write_json(fields, args.out_json)
    write_csv(fields, args.out_csv)

    # Print the questions you will ask the user (company-side only)
    questions = [f"- {x.label_for_user}: {x.how_to_fill}" for x in fields if x.ask_user]
    print("\nQuestions to ask user:")
    print("\n".join(questions))

    print(f"\nWrote:\n- {args.out_json}\n- {args.out_csv}")


if __name__ == "__main__":
    # Ensure key is present for LLM mode
    # export OPENAI_API_KEY="..."
    main()